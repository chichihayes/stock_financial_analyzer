"""
Core financial-analysis logic, extracted from the original Streamlit app
(finance_app.py) with all Streamlit calls removed so it can run inside
stateless Vercel Python serverless functions.

Config that used to live in st.secrets now comes from environment
variables (set OPENROUTER_API_KEY in your Vercel project settings).
"""
import math
import os
import time

import numpy as np
import pandas as pd
import requests
from docx import Document
from io import BytesIO
from yahooquery import Ticker


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def format_number(value):
    """Format numbers for better readability."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "N/A"
    if isinstance(value, (int, float)):
        if abs(value) >= 1e12:
            return f"${value / 1e12:.2f}T"
        elif abs(value) >= 1e9:
            return f"${value / 1e9:.2f}B"
        elif abs(value) >= 1e6:
            return f"${value / 1e6:.2f}M"
        elif abs(value) >= 1e3:
            return f"${value / 1e3:.2f}K"
        else:
            return f"${value:.2f}" if value > 1 else f"{value:.3f}"
    return str(value)


def format_percentage(value):
    """Format ratio as percentage."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "N/A"
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%" if abs(value) < 10 else f"{value:.2f}%"
    return str(value)


# ---------------------------------------------------------------------------
# Data retrieval
# ---------------------------------------------------------------------------

def clean_income_statement(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean income statement data."""
    ticker = Ticker(ticker_symbol)
    income_statement = ticker.income_statement(frequency="a")

    numeric_columns = income_statement.select_dtypes(include=["float64", "int64"]).columns
    income_statement_cleaned = income_statement[numeric_columns]

    income_statement_latest = income_statement_cleaned.iloc[-1:].T.reset_index()
    income_statement_latest.columns = ["parameters", "value"]

    return income_statement_latest


def clean_balance_sheet(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean balance sheet data."""
    ticker = Ticker(ticker_symbol)
    balance_sheet = ticker.balance_sheet(frequency="a")

    numeric_columns = balance_sheet.select_dtypes(include=["float64", "int64"]).columns
    balance_sheet_cleaned = balance_sheet[numeric_columns]

    balance_sheet_latest = balance_sheet_cleaned.iloc[-1:].T.reset_index()
    balance_sheet_latest.columns = ["parameters", "value"]

    return balance_sheet_latest


def clean_cash_flow_statement(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean cash flow statement data."""
    ticker = Ticker(ticker_symbol)
    cash_flow = ticker.cash_flow(frequency="a").T

    cash_flow_cleaned = cash_flow.drop(index=["symbol", "periodType", "currencyCode"], errors="ignore")

    cash_flow_cleaned.loc["asOfDate"] = cash_flow_cleaned.loc["asOfDate"].astype(str).str.replace("_", "-")
    new_column_names = pd.to_datetime(cash_flow_cleaned.loc["asOfDate"]).dt.strftime("%Y_%m_%d")
    cash_flow_cleaned.columns = new_column_names

    cash_flow_cleaned = cash_flow_cleaned.drop(index="asOfDate")

    latest_year = cash_flow_cleaned.columns[-1]
    cash_flow_latest = cash_flow_cleaned[[latest_year]]

    cash_flow_latest.index.name = "parameters"
    cash_flow_latest.reset_index(inplace=True)

    return cash_flow_latest


# ---------------------------------------------------------------------------
# Ratio calculation
# ---------------------------------------------------------------------------

def calculate_financial_ratios(balance_sheet: pd.DataFrame, income_statement: pd.DataFrame,
                                cash_flow: pd.DataFrame) -> dict:
    """Calculate comprehensive financial ratios."""
    try:
        balance_latest = balance_sheet.columns[-1]
        income_latest = income_statement.columns[-1]
        cash_flow_latest = cash_flow.columns[-1]

        balance_dict = balance_sheet.set_index("parameters")[balance_latest].to_dict()
        income_dict = income_statement.set_index("parameters")[income_latest].to_dict()
        cash_flow_dict = cash_flow.set_index("parameters")[cash_flow_latest].to_dict()
    except Exception:
        balance_dict = balance_sheet.set_index("parameters")["value"].to_dict() if "value" in balance_sheet.columns else {}
        income_dict = income_statement.set_index("parameters")["value"].to_dict() if "value" in income_statement.columns else {}
        cash_flow_dict = cash_flow.set_index("parameters")["value"].to_dict() if "value" in cash_flow.columns else {}
        balance_latest = income_latest = cash_flow_latest = "Latest"

    def get_value(dictionary, keys):
        for key in keys:
            value = dictionary.get(key)
            if value is not None and not pd.isna(value):
                return value
        return None

    current_assets = get_value(balance_dict, ["CurrentAssets", "currentAssets", "current_assets"])
    current_liabilities = get_value(balance_dict, ["CurrentLiabilities", "currentLiabilities"])
    cash_equivalents = get_value(balance_dict, ["CashAndCashEquivalents", "cashAndEquivalents"])
    inventory = get_value(balance_dict, ["Inventory", "inventory"])
    total_assets = get_value(balance_dict, ["TotalAssets", "totalAssets"])
    total_liabilities = get_value(balance_dict, ["TotalLiabilitiesNetMinorityInterest", "totalLiabilities"])
    total_equity = get_value(balance_dict, ["TotalEquityGrossMinorityInterest", "totalEquity"])
    ordinary_shares = get_value(balance_dict, ["OrdinarySharesNumber", "sharesOutstanding"])

    net_income = get_value(income_dict, ["NetIncome", "netIncome", "net_income"])
    revenue = get_value(income_dict, ["TotalRevenue", "totalRevenue", "revenue"])
    eps = get_value(income_dict, ["DilutedEPS", "dilutedEps", "eps"])

    free_cash_flow = get_value(cash_flow_dict, ["FreeCashFlow", "freeCashFlow"])

    ratios = {
        "Balance Sheet Date": balance_latest,
        "Income Statement Date": income_latest,
        "Cash Flow Date": cash_flow_latest,
        "Current Ratio": (current_assets / current_liabilities) if current_assets and current_liabilities else None,
        "Quick Ratio": ((current_assets - (inventory or 0)) / current_liabilities) if current_assets and current_liabilities else None,
        "Cash Ratio": (cash_equivalents / current_liabilities) if cash_equivalents and current_liabilities else None,
        "Debt-to-Equity Ratio": (total_liabilities / total_equity) if total_liabilities and total_equity else None,
        "Debt-to-Assets Ratio": (total_liabilities / total_assets) if total_liabilities and total_assets else None,
        "Book Value per Share (BVPS)": (total_equity / ordinary_shares) if total_equity and ordinary_shares else None,
        "Return on Equity (ROE)": (net_income / total_equity) if net_income and total_equity else None,
        "Return on Assets (ROA)": (net_income / total_assets) if net_income and total_assets else None,
        "Net Margin": (net_income / revenue) if net_income and revenue else None,
        "EPS": eps,
        "Revenue": revenue,
        "Net Income": net_income,
        "Free Cash Flow": free_cash_flow,
    }

    return ratios


def to_jsonable(ratios: dict) -> dict:
    """Convert numpy/pandas scalar types (and NaN) to plain JSON-safe values."""
    clean = {}
    for key, value in ratios.items():
        if value is None:
            clean[key] = None
        elif isinstance(value, (pd.Timestamp,)):
            clean[key] = str(value)
        elif isinstance(value, (np.floating, float)):
            clean[key] = None if (isinstance(value, float) and math.isnan(value)) else float(value)
        elif isinstance(value, (np.integer,)):
            clean[key] = int(value)
        else:
            clean[key] = value
    return clean


def from_jsonable(ratios: dict) -> dict:
    """Inverse-ish helper: just returns the dict as-is (values already plain types).

    Kept as a separate name for readability at call sites that receive ratios
    back from the client (e.g. the /api/report endpoint).
    """
    return dict(ratios)


# ---------------------------------------------------------------------------
# Analysis generation
# ---------------------------------------------------------------------------

def generate_financial_analysis_openrouter(ticker_symbol, ratios):
    """Generate financial analysis using OpenRouter. Returns (analysis, log_lines)."""

    API_URL = "https://openrouter.ai/api/v1"
    log = []

    prompt = f"""
    Analyze the following financial ratios for {ticker_symbol}:

    - Current Ratio: {format_number(ratios['Current Ratio'])}
    - Quick Ratio: {format_number(ratios['Quick Ratio'])}
    - Debt-to-Equity: {format_number(ratios['Debt-to-Equity Ratio'])}
    - ROE: {format_percentage(ratios['Return on Equity (ROE)'])}
    - ROA: {format_percentage(ratios['Return on Assets (ROA)'])}
    - Net Margin: {format_percentage(ratios['Net Margin'])}

    Based on this data, please provide an analysis that clearly explains:
      1. The strengths and weaknesses of the company's financial position.
      2. Reasons why investors should consider buying the stock.
      3. Reasons why investors might decide not to buy the stock.

    Provide your analysis in bullet points (at least 5 points), detailing your rationale regarding liquidity, profitability, and risk.
    Do not include the prompt text in your response. Respond only with your analysis.
    """

    start_time = time.time()

    try:
        api_key = os.environ.get("OPENROUTER_API_KEY")
        if not api_key:
            log.append("No OpenRouter API key provided. Falling back to rule-based analysis.")
            return generate_rule_based_analysis(ratios), log

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        data = {
            "model": "google/gemini-2.5-flash-lite",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
        }

        log.append("Sending request to OpenRouter API...")
        response = requests.post(API_URL + "/chat/completions", json=data, headers=headers, timeout=30)
        log.append(f"API Status Code: {response.status_code}")

        if response.status_code == 200:
            analysis = response.json()["choices"][0]["message"]["content"].strip()
            end_time = time.time()
            log.append(f"Analysis generated in {end_time - start_time:.2f} seconds")
            return analysis, log
        else:
            log.append(f"API Response Error: {response.status_code}")
            log.append(f"Response Details: {response.text[:100]}...")
            return generate_rule_based_analysis(ratios), log

    except Exception as e:
        log.append(f"Error connecting to OpenRouter API: {str(e)}")
        return generate_rule_based_analysis(ratios), log


def generate_rule_based_analysis(ratios):
    """Generate a basic analysis based on rules when API is unavailable."""
    analysis = "Financial Analysis:\n\n"

    current_ratio = ratios.get("Current Ratio")
    if current_ratio:
        if current_ratio > 2:
            analysis += "• Liquidity is strong with a current ratio of {:.2f}x, suggesting the company can easily meet short-term obligations.\n".format(current_ratio)
        elif current_ratio > 1:
            analysis += "• Liquidity is adequate with a current ratio of {:.2f}x, sufficient to cover short-term liabilities.\n".format(current_ratio)
        else:
            analysis += "• Liquidity is concerning with a current ratio of {:.2f}x, indicating potential challenges meeting short-term obligations.\n".format(current_ratio)

    debt_equity = ratios.get("Debt-to-Equity Ratio")
    if debt_equity:
        if debt_equity > 2:
            analysis += "• The debt-to-equity ratio of {:.2f}x is high, suggesting significant leverage and financial risk.\n".format(debt_equity)
        elif debt_equity > 1:
            analysis += "• The debt-to-equity ratio of {:.2f}x indicates moderate leverage, with debt exceeding equity.\n".format(debt_equity)
        else:
            analysis += "• The debt-to-equity ratio of {:.2f}x is conservative, indicating lower financial risk.\n".format(debt_equity)

    roe = ratios.get("Return on Equity (ROE)")
    if roe:
        if roe > 0.15:
            analysis += "• Return on Equity (ROE) of {:.1f}% is excellent, indicating efficient use of shareholder capital.\n".format(roe * 100)
        elif roe > 0.08:
            analysis += "• Return on Equity (ROE) of {:.1f}% is solid, showing reasonable returns on shareholder investments.\n".format(roe * 100)
        else:
            analysis += "• Return on Equity (ROE) of {:.1f}% could be improved to enhance shareholder returns.\n".format(roe * 100)

    net_margin = ratios.get("Net Margin")
    if net_margin:
        if net_margin > 0.15:
            analysis += "• Net profit margin of {:.1f}% is strong, demonstrating effective cost management and pricing power.\n".format(net_margin * 100)
        elif net_margin > 0.05:
            analysis += "• Net profit margin of {:.1f}% is acceptable but could be improved through cost optimization.\n".format(net_margin * 100)
        else:
            analysis += "• Net profit margin of {:.1f}% is low, suggesting the need for revenue growth or cost reduction strategies.\n".format(net_margin * 100)

    analysis += "\nRecommendations:\n"
    analysis += "• Review the complete financial statements for a more comprehensive understanding of the company's position.\n"
    analysis += "• Compare these metrics with industry peers to gain competitive context.\n"
    analysis += "• Monitor trends over time to identify improvement or deterioration in financial health.\n"
    analysis += "• Consider the company's growth stage and industry norms when evaluating these metrics.\n"
    analysis += "• Consult with a financial advisor for personalized investment advice based on this analysis.\n"

    return analysis


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

def create_docx_report(ticker_symbol, ratios, analysis) -> BytesIO:
    """Create a Word document report with financial analysis."""
    doc = Document()

    doc.add_heading(f"Financial Analysis Report: {ticker_symbol}", 0)

    doc.add_heading("Statement Dates", level=1)
    doc.add_paragraph(f"Balance Sheet: {ratios['Balance Sheet Date']}")
    doc.add_paragraph(f"Income Statement: {ratios['Income Statement Date']}")
    doc.add_paragraph(f"Cash Flow: {ratios['Cash Flow Date']}")

    doc.add_heading("Financial Performance", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"

    metrics = [
        ("Revenue", format_number(ratios["Revenue"])),
        ("Net Income", format_number(ratios["Net Income"])),
        ("EPS", format_number(ratios["EPS"])),
        ("Free Cash Flow", format_number(ratios["Free Cash Flow"])),
        ("ROE", format_percentage(ratios["Return on Equity (ROE)"])),
        ("ROA", format_percentage(ratios["Return on Assets (ROA)"])),
        ("Net Margin", format_percentage(ratios["Net Margin"])),
        ("Current Ratio", f"{ratios['Current Ratio']:.2f}x" if ratios["Current Ratio"] else "N/A"),
        ("Quick Ratio", f"{ratios['Quick Ratio']:.2f}x" if ratios["Quick Ratio"] else "N/A"),
        ("Cash Ratio", f"{ratios['Cash Ratio']:.2f}x" if ratios["Cash Ratio"] else "N/A"),
        ("Debt-to-Equity", f"{ratios['Debt-to-Equity Ratio']:.2f}x" if ratios["Debt-to-Equity Ratio"] else "N/A"),
        ("Debt-to-Assets", f"{ratios['Debt-to-Assets Ratio']:.2f}x" if ratios["Debt-to-Assets Ratio"] else "N/A"),
        ("Book Value per Share", format_number(ratios["Book Value per Share (BVPS)"])),
    ]

    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)

    doc.add_heading("Financial Analysis", level=1)
    doc.add_paragraph(analysis)

    doc.add_paragraph("Generated by Orbann_ai")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
