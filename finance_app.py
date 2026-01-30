import streamlit as st
import pandas as pd
import numpy as np
from yahooquery import Ticker
import os
import time
import requests
from docx import Document
from io import BytesIO

def format_number(value):
    """Format numbers for better readability"""
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, (int, float)):
        if abs(value) >= 1e12:
            return f"${value/1e12:.2f}T"
        elif abs(value) >= 1e9:
            return f"${value/1e9:.2f}B"
        elif abs(value) >= 1e6:
            return f"${value/1e6:.2f}M"
        elif abs(value) >= 1e3:
            return f"${value/1e3:.2f}K"
        else:
            return f"${value:.2f}" if value > 1 else f"{value:.3f}"
    return str(value)

def format_percentage(value):
    """Format ratio as percentage"""
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%" if abs(value) < 10 else f"{value:.2f}%"
    return str(value)

def clean_income_statement(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean income statement data."""
    ticker = Ticker(ticker_symbol)
    income_statement = ticker.income_statement(frequency="a")

    # Keep only the numeric data columns
    numeric_columns = income_statement.select_dtypes(include=['float64', 'int64']).columns
    income_statement_cleaned = income_statement[numeric_columns]

    # Get the most recent period's data
    income_statement_latest = income_statement_cleaned.iloc[-1:].T.reset_index()
    income_statement_latest.columns = ['parameters', 'value']

    return income_statement_latest

def clean_balance_sheet(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean balance sheet data."""
    ticker = Ticker(ticker_symbol)
    balance_sheet = ticker.balance_sheet(frequency="a")

    # Keep only the numeric data columns
    numeric_columns = balance_sheet.select_dtypes(include=['float64', 'int64']).columns
    balance_sheet_cleaned = balance_sheet[numeric_columns]

    # Get the most recent period's data
    balance_sheet_latest = balance_sheet_cleaned.iloc[-1:].T.reset_index()
    balance_sheet_latest.columns = ['parameters', 'value']

    return balance_sheet_latest

def clean_cash_flow_statement(ticker_symbol: str) -> pd.DataFrame:
    """Fetch and clean cash flow statement data."""
    ticker = Ticker(ticker_symbol)
    cash_flow = ticker.cash_flow(frequency="a").T

    # Drop unnecessary rows
    cash_flow_cleaned = cash_flow.drop(index=["symbol", "periodType", "currencyCode"], errors="ignore")

    # Fix date parsing issue
    cash_flow_cleaned.loc["asOfDate"] = cash_flow_cleaned.loc["asOfDate"].astype(str).str.replace("_", "-")
    new_column_names = pd.to_datetime(cash_flow_cleaned.loc["asOfDate"]).dt.strftime("%Y_%m_%d")
    cash_flow_cleaned.columns = new_column_names

    # Drop 'asOfDate'
    cash_flow_cleaned = cash_flow_cleaned.drop(index="asOfDate")

    # Extract latest year data
    latest_year = cash_flow_cleaned.columns[-1]
    cash_flow_latest = cash_flow_cleaned[[latest_year]]

    # Format DataFrame
    cash_flow_latest.index.name = 'parameters'
    cash_flow_latest.reset_index(inplace=True)

    return cash_flow_latest

def calculate_financial_ratios(balance_sheet: pd.DataFrame, income_statement: pd.DataFrame, cash_flow: pd.DataFrame) -> dict:
    """Calculate comprehensive financial ratios."""
    try:
        # Get the latest column names from each DataFrame
        balance_latest = balance_sheet.columns[-1]
        income_latest = income_statement.columns[-1]
        cash_flow_latest = cash_flow.columns[-1]
        
        # Convert DataFrames to dicts for easier access
        balance_dict = balance_sheet.set_index("parameters")[balance_latest].to_dict()
        income_dict = income_statement.set_index("parameters")[income_latest].to_dict()
        cash_flow_dict = cash_flow.set_index("parameters")[cash_flow_latest].to_dict()
    except:
        # Alternative method if the above fails
        balance_dict = balance_sheet.set_index("parameters")['value'].to_dict() if 'value' in balance_sheet.columns else {}
        income_dict = income_statement.set_index("parameters")['value'].to_dict() if 'value' in income_statement.columns else {}
        cash_flow_dict = cash_flow.set_index("parameters")['value'].to_dict() if 'value' in cash_flow.columns else {}
        balance_latest = income_latest = cash_flow_latest = "Latest"

    # Function to get value with multiple possible keys
    def get_value(dictionary, keys):
        """Try multiple keys to get a value"""
        for key in keys:
            value = dictionary.get(key)
            if value is not None and not pd.isna(value):
                return value
        return None

    # Extract metrics with fallback keys
    current_assets = get_value(balance_dict, ["CurrentAssets", "currentAssets", "current_assets"])
    current_liabilities = get_value(balance_dict, ["CurrentLiabilities", "currentLiabilities"])
    cash_equivalents = get_value(balance_dict, ["CashAndCashEquivalents", "cashAndEquivalents"])
    inventory = get_value(balance_dict, ["Inventory", "inventory"])
    total_assets = get_value(balance_dict, ["TotalAssets", "totalAssets"])
    total_liabilities = get_value(balance_dict, ["TotalLiabilitiesNetMinorityInterest", "totalLiabilities"])
    total_equity = get_value(balance_dict, ["TotalEquityGrossMinorityInterest", "totalEquity"])
    ordinary_shares = get_value(balance_dict, ["OrdinarySharesNumber", "sharesOutstanding"])

    # Income statement metrics with alternative keys
    net_income = get_value(income_dict, ["NetIncome", "netIncome", "net_income"])
    revenue = get_value(income_dict, ["TotalRevenue", "totalRevenue", "revenue"])
    eps = get_value(income_dict, ["DilutedEPS", "dilutedEps", "eps"])
    ebitda = get_value(income_dict, ["EBITDA", "ebitda"])

    # Cash flow metrics
    free_cash_flow = get_value(cash_flow_dict, ["FreeCashFlow", "freeCashFlow"])

    # Calculate ratios
    ratios = {
        "Balance Sheet Date": balance_latest,
        "Income Statement Date": income_latest,
        "Cash Flow Date": cash_flow_latest,
        "Current Ratio": (current_assets / current_liabilities) if current_assets and current_liabilities else None,
        "Quick Ratio": ((current_assets - (inventory or 0)) / current_liabilities) if current_assets and current_liabilities else None,
        "Cash Ratio": (cash_equivalents / current_liabilities) if cash_equivalents and current_liabilities else None,
        "Debt-to-Equity Ratio": (total_liabilities / total_equity) if total_liabilities and total_equity else None,
        "Debt-to-Assets Ratio": (total_liabilities / total_assets) if total_liabilities and total_assets else None,
        "Book Value per Share (BVPS)": (total_equity / ordinary_shares) if total_equity and ordinary_shares else None,
        "Return on Equity (ROE)": (net_income / total_equity) if net_income and total_equity else None,
        "Return on Assets (ROA)": (net_income / total_assets) if net_income and total_assets else None,
        "Net Margin": (net_income / revenue) if net_income and revenue else None,
        "EPS": eps,
        "Revenue": revenue,
        "Net Income": net_income,
        "Free Cash Flow": free_cash_flow
    }
    
    return ratios

import requests
import time
import streamlit as st

def generate_financial_analysis_openrouter(ticker_symbol, ratios):
    """Generate financial analysis using OpenRouter's DeepSeek-R1:free"""

    API_URL = "https://openrouter.ai/api/v1"
    
    prompt = f"""
    Analyze the following financial ratios for {ticker_symbol}:

    - Current Ratio: {format_number(ratios['Current Ratio'])}
    - Quick Ratio: {format_number(ratios['Quick Ratio'])}
    - Debt-to-Equity: {format_number(ratios['Debt-to-Equity Ratio'])}
    - ROE: {format_percentage(ratios['Return on Equity (ROE)'])}
    - ROA: {format_percentage(ratios['Return on Assets (ROA)'])}
    - Net Margin: {format_percentage(ratios['Net Margin'])}

    Based on this data, please provide an analysis that clearly explains:
      1. The strengths and weaknesses of the company's financial position.
      2. Reasons why investors should consider buying the stock.
      3. Reasons why investors might decide not to buy the stock.

    Provide your analysis in bullet points (at least 5 points), detailing your rationale regarding liquidity, profitability, and risk.
    Do not include the prompt text in your response. Respond only with your analysis.
    """

    start_time = time.time()

    try:
        # Retrieve the OpenRouter API key from secrets
        api_key = st.secrets.get("general", {}).get("OPENROUTER_API_KEY")
        if not api_key:
            st.sidebar.error("No OpenRouter API key provided. Please set it in your secrets.")
            return generate_rule_based_analysis(ratios)

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "google/gemini-2.5-flash-lite",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }

        st.sidebar.info("Sending request to OpenRouter API using DeepSeek-R1...")
        response = requests.post(API_URL + "/chat/completions", json=data, headers=headers, timeout=30)

        st.sidebar.info(f"API Status Code: {response.status_code}")

        if response.status_code == 200:
            # Extract the generated text
            analysis = response.json()["choices"][0]["message"]["content"].strip()
            
            end_time = time.time()
            st.sidebar.success(f"Analysis generated in {end_time - start_time:.2f} seconds")
            return analysis
        else:
            st.sidebar.warning(f"API Response Error: {response.status_code}")
            st.sidebar.warning(f"Response Details: {response.text[:100]}...")
            return generate_rule_based_analysis(ratios)

    except Exception as e:
        st.sidebar.error(f"Error connecting to OpenRouter API: {str(e)}")
        return generate_rule_based_analysis(ratios)


def generate_rule_based_analysis(ratios):
    """Generate a basic analysis based on rules when API is unavailable"""
    analysis = f"Financial Analysis:\n\n"
    
    # Liquidity analysis
    current_ratio = ratios.get('Current Ratio')
    if current_ratio:
        if current_ratio > 2:
            analysis += "• Liquidity is strong with a current ratio of {:.2f}x, suggesting the company can easily meet short-term obligations.\n".format(current_ratio)
        elif current_ratio > 1:
            analysis += "• Liquidity is adequate with a current ratio of {:.2f}x, sufficient to cover short-term liabilities.\n".format(current_ratio)
        else:
            analysis += "• Liquidity is concerning with a current ratio of {:.2f}x, indicating potential challenges meeting short-term obligations.\n".format(current_ratio)
    
    # Solvency analysis
    debt_equity = ratios.get('Debt-to-Equity Ratio')
    if debt_equity:
        if debt_equity > 2:
            analysis += "• The debt-to-equity ratio of {:.2f}x is high, suggesting significant leverage and financial risk.\n".format(debt_equity)
        elif debt_equity > 1:
            analysis += "• The debt-to-equity ratio of {:.2f}x indicates moderate leverage, with debt exceeding equity.\n".format(debt_equity)
        else:
            analysis += "• The debt-to-equity ratio of {:.2f}x is conservative, indicating lower financial risk.\n".format(debt_equity)
    
    # Profitability analysis
    roe = ratios.get('Return on Equity (ROE)')
    if roe:
        if roe > 0.15:
            analysis += "• Return on Equity (ROE) of {:.1f}% is excellent, indicating efficient use of shareholder capital.\n".format(roe * 100)
        elif roe > 0.08:
            analysis += "• Return on Equity (ROE) of {:.1f}% is solid, showing reasonable returns on shareholder investments.\n".format(roe * 100)
        else:
            analysis += "• Return on Equity (ROE) of {:.1f}% could be improved to enhance shareholder returns.\n".format(roe * 100)
    
    # Net margin analysis
    net_margin = ratios.get('Net Margin')
    if net_margin:
        if net_margin > 0.15:
            analysis += "• Net profit margin of {:.1f}% is strong, demonstrating effective cost management and pricing power.\n".format(net_margin * 100)
        elif net_margin > 0.05:
            analysis += "• Net profit margin of {:.1f}% is acceptable but could be improved through cost optimization.\n".format(net_margin * 100)
        else:
            analysis += "• Net profit margin of {:.1f}% is low, suggesting the need for revenue growth or cost reduction strategies.\n".format(net_margin * 100)
    
    # Add recommendations
    analysis += "\nRecommendations:\n"
    analysis += "• Review the complete financial statements for a more comprehensive understanding of the company's position.\n"
    analysis += "• Compare these metrics with industry peers to gain competitive context.\n"
    analysis += "• Monitor trends over time to identify improvement or deterioration in financial health.\n"
    analysis += "• Consider the company's growth stage and industry norms when evaluating these metrics.\n"
    analysis += "• Consult with a financial advisor for personalized investment advice based on this analysis.\n"
    
    return analysis

def create_docx_report(ticker_symbol, ratios, analysis):
    """Create a Word document report with financial analysis"""
    doc = Document()
    
    # Add title
    doc.add_heading(f'Financial Analysis Report: {ticker_symbol}', 0)
    
    # Add date information
    doc.add_heading('Statement Dates', level=1)
    doc.add_paragraph(f"Balance Sheet: {ratios['Balance Sheet Date']}")
    doc.add_paragraph(f"Income Statement: {ratios['Income Statement Date']}")
    doc.add_paragraph(f"Cash Flow: {ratios['Cash Flow Date']}")
    
    # Add financial metrics
    doc.add_heading('Financial Performance', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Add metrics to table
    metrics = [
        ("Revenue", format_number(ratios["Revenue"])),
        ("Net Income", format_number(ratios["Net Income"])),
        ("EPS", format_number(ratios["EPS"])),
        ("Free Cash Flow", format_number(ratios["Free Cash Flow"])),
        ("ROE", format_percentage(ratios["Return on Equity (ROE)"])),
        ("ROA", format_percentage(ratios["Return on Assets (ROA)"])),
        ("Net Margin", format_percentage(ratios["Net Margin"])),
        ("Current Ratio", f"{ratios['Current Ratio']:.2f}x" if ratios['Current Ratio'] else "N/A"),
        ("Quick Ratio", f"{ratios['Quick Ratio']:.2f}x" if ratios['Quick Ratio'] else "N/A"),
        ("Cash Ratio", f"{ratios['Cash Ratio']:.2f}x" if ratios['Cash Ratio'] else "N/A"),
        ("Debt-to-Equity", f"{ratios['Debt-to-Equity Ratio']:.2f}x" if ratios['Debt-to-Equity Ratio'] else "N/A"),
        ("Debt-to-Assets", f"{ratios['Debt-to-Assets Ratio']:.2f}x" if ratios['Debt-to-Assets Ratio'] else "N/A"),
        ("Book Value per Share", format_number(ratios["Book Value per Share (BVPS)"]))
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    
    # Add analysis
    doc.add_heading('Financial Analysis', level=1)
    doc.add_paragraph(analysis)
    
    # Add footer
    doc.add_paragraph('Generated by Orbann_ai')
    
    # Save to BytesIO object to allow downloading
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def main():
    st.set_page_config(page_title="Financial Analysis Dashboard", layout="wide")
    
    # Add a debug mode toggle
    debug_mode = False
    if 'debug_mode' not in st.session_state:
        st.session_state.debug_mode = debug_mode
    
    # Title and Description
    st.title("Orbann_ai")
    st.markdown("""
    A comprehensive tool for financial statement analysis
    
    Contact: [gizhayes27@gmail.com](mailto:gizhayes27@gmail.com) | 
    Community: [Join our Telegram](https://t.me/+0WciZpJaSOhhMmM0)
    """)
    st.markdown("---")

    # Sidebar settings
    st.sidebar.title("Settings")
    use_ai = st.sidebar.checkbox("Generate Analysis", value=True)
    
    # Add debug mode option to the sidebar
    st.session_state.debug_mode = st.sidebar.checkbox("Debug Mode", value=st.session_state.debug_mode)
    
    # Display API configuration in debug mode
    if st.session_state.debug_mode:
        st.sidebar.subheader("API Configuration")
        api_key_display = st.secrets.get("general", {}).get("OPENROUTER_API_KEY", 'Using default key')
        st.sidebar.text(f"API Key: {api_key_display[:5]}..." if len(api_key_display) > 10 else "No OpenRouter API key found")
    
    # Use columns for better space utilization
    left_column, right_column = st.columns([1, 3])

    # Store ratios and analysis in session state
    if 'ratios' not in st.session_state:
        st.session_state.ratios = None
    if 'analysis' not in st.session_state:
        st.session_state.analysis = None

    # Sidebar content moves to left column
    with left_column:
        st.header("Input Parameters")
        ticker_symbol = st.text_input("Enter Stock Ticker:", value="AAPL").upper()
        if st.button("Analyze", use_container_width=True):
            try:
                with st.spinner("Fetching financial data..."):
                    balance_sheet = clean_balance_sheet(ticker_symbol)
                    income_statement = clean_income_statement(ticker_symbol)
                    cash_flow = clean_cash_flow_statement(ticker_symbol)
                    st.session_state.ratios = calculate_financial_ratios(balance_sheet, income_statement, cash_flow)
                
                if use_ai:
                    with st.spinner("Generating analysis..."):
                        st.session_state.analysis = generate_financial_analysis_openrouter(
                            ticker_symbol, 
                            st.session_state.ratios
                        )
                else:
                    # Use rule-based analysis if generation is disabled
                    st.session_state.analysis = generate_rule_based_analysis(
                        st.session_state.ratios
                    )
                
                st.success(f"Analysis completed for {ticker_symbol}")
            except Exception as e:
                st.error(f"Error occurred: {str(e)}")
                st.warning("Please check the ticker symbol and try again.")
                
                # Show detailed error in debug mode
                if st.session_state.debug_mode:
                    import traceback
                    st.error(traceback.format_exc())

    # Main content area in right column
    with right_column:
        if st.session_state.ratios is not None:
            ratios = st.session_state.ratios
            
            # Dates Information in a more compact format
            st.markdown("### Statement Dates")
            dates_cols = st.columns(3)
            with dates_cols[0]:
                st.info(f"Balance Sheet: {ratios['Balance Sheet Date']}")
            with dates_cols[1]:
                st.info(f"Income Statement: {ratios['Income Statement Date']}")
            with dates_cols[2]:
                st.info(f"Cash Flow: {ratios['Cash Flow Date']}")

            # Financial Performance with better formatting
            st.markdown("### Financial Performance")
            perf_cols = st.columns(4)
            performance_metrics = [
                ("Revenue", ratios["Revenue"]),
                ("Net Income", ratios["Net Income"]),
                ("EPS", ratios["EPS"]),
                ("Free Cash Flow", ratios["Free Cash Flow"])
            ]
            for col, (label, value) in zip(perf_cols, performance_metrics):
                with col:
                    st.metric(label, format_number(value))

            # Profitability Metrics
            st.markdown("### Profitability Metrics")
            profit_cols = st.columns(3)
            profitability_metrics = [
                ("ROE", ratios["Return on Equity (ROE)"]),
                ("ROA", ratios["Return on Assets (ROA)"]),
                ("Net Margin", ratios["Net Margin"])
            ]
            for col, (label, value) in zip(profit_cols, profitability_metrics):
                with col:
                    st.metric(label, format_percentage(value))

            # Liquidity & Solvency Combined
            st.markdown("### Liquidity & Solvency")
            ratio_cols = st.columns(3)
            ratio_metrics = [
                ("Current Ratio", ratios["Current Ratio"]),
                ("Quick Ratio", ratios["Quick Ratio"]),
                ("Cash Ratio", ratios["Cash Ratio"])
            ]
            for col, (label, value) in zip(ratio_cols, ratio_metrics):
                with col:
                    st.metric(label, f"{value:.2f}x" if value else "N/A")
            
            # Additional row for Solvency
            solvency_cols = st.columns(3)
            solvency_metrics = [
                ("Debt/Equity", ratios["Debt-to-Equity Ratio"]),
                ("Debt/Assets", ratios["Debt-to-Assets Ratio"]),
                ("Book Value/Share", ratios["Book Value per Share (BVPS)"])
            ]
            for col, (label, value) in zip(solvency_cols, solvency_metrics):
                with col:
                    formatted_value = f"{value:.2f}x" if "Debt" in label and value else format_number(value)
                    st.metric(label, formatted_value)
            
            # Analysis section
            if st.session_state.analysis:
                st.markdown("### Financial Analysis")
                st.write(st.session_state.analysis)
            
            # Download button for Word document
            if st.session_state.ratios:
                report = create_docx_report(
                    ticker_symbol, 
                    st.session_state.ratios, 
                    st.session_state.analysis if st.session_state.analysis else "No analysis generated."
                )
                st.download_button(
                    label="Download Full Report (DOCX)",
                    data=report,
                    file_name=f"{ticker_symbol}_Financial_Analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            # Footnote
            st.markdown("---")
            st.header("Support")
            st.markdown("""
                        Need help or you want to support our startup?
                        - Email: [gizhayes27@gmail.com](mailto:gizhayes27@gmail.com)
                        - Join our [Telegram Group](https://t.me/+0WciZpJaSOhhMmM0)""")
            st.caption("orbann_ai")

if __name__ == "__main__":
    main()
